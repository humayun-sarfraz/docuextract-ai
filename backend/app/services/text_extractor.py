import logging
from pathlib import Path

logger = logging.getLogger("docuextract.text")


def extract_text_from_file(file_path: str) -> str:
    path = Path(file_path)
    ext = path.suffix.lower()
    logger.info("Extracting text from %s (%s)", path.name, ext)

    if ext == ".txt":
        return _extract_txt(path)
    elif ext == ".pdf":
        return _extract_pdf(path)
    elif ext == ".docx":
        return _extract_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _extract_pdf(path: Path) -> str:
    try:
        import pymupdf
    except ImportError:
        raise ImportError("pymupdf is required for PDF support. Install it: pip install pymupdf")

    text_parts = []
    with pymupdf.open(str(path)) as doc:
        for page in doc:
            text_parts.append(page.get_text())
    text = "\n".join(text_parts)
    logger.info("PDF extracted: %d pages, %d chars", len(text_parts), len(text))
    return text


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX support. Install it: pip install python-docx")

    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs)
    logger.info("DOCX extracted: %d paragraphs, %d chars", len(doc.paragraphs), len(text))
    return text
